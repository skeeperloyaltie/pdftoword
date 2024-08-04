import os
from pdf2docx import Converter
from docx import Document
import fitz  # PyMuPDF
import traceback

def convert_pdf_to_word(pdf_folder_path, word_folder_path):
    # Ensure the output directory exists
    if not os.path.exists(word_folder_path):
        os.makedirs(word_folder_path)

    # Iterate through all PDF files in the specified folder
    for filename in os.listdir(pdf_folder_path):
        if filename.endswith(".pdf"):
            pdf_file_path = os.path.join(pdf_folder_path, filename)
            word_file_path = os.path.join(word_folder_path, f"{os.path.splitext(filename)[0]}.docx")
            
            try:
                # Convert PDF to Word using pdf2docx
                cv = Converter(pdf_file_path)
                cv.convert(word_file_path, start=0, end=None)
                cv.close()
                print(f"Converted {filename} using pdf2docx")

                # Ensure all texts are editable
                doc = Document(word_file_path)
                doc.save(word_file_path)
                
            except Exception as e:
                print(f"Failed to convert {filename} using pdf2docx, attempting with PyMuPDF")
                print(traceback.format_exc())
                try:
                    # Fallback conversion using PyMuPDF
                    doc = fitz.open(pdf_file_path)
                    text = ""
                    for page_num in range(doc.page_count):
                        page = doc.load_page(page_num)
                        text += page.get_text()

                    word_doc = Document()
                    word_doc.add_paragraph(text)
                    word_doc.save(word_file_path)
                    print(f"Converted {filename} using PyMuPDF")
                except Exception as e:
                    print(f"Failed to convert {filename} using PyMuPDF")
                    print(traceback.format_exc())

    print(f"All PDF files in {pdf_folder_path} have been processed")

# Get user input for folder path
pdf_folder = input("Enter the path to the folder containing PDF files: ").strip()
if pdf_folder.startswith("~"):
    pdf_folder = os.path.expanduser(pdf_folder)

word_folder = pdf_folder  # Save Word files in the same folder

convert_pdf_to_word(pdf_folder, word_folder)
